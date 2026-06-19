import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from quantlab.reports import (
    WORD_REPORT_TYPES,
    artifact_counts,
    cleanup_report_junk,
    collect_report_artifacts,
    experiment_summaries_frame,
    export_experiment_docx,
    filter_report_artifacts_frame,
    latest_report_artifact,
    load_experiment_detail,
    report_activity_frame,
    report_artifacts_frame,
    report_dashboard_cards,
    run_status_counts_frame,
    run_metadata_summaries_frame,
    search_report_artifacts_frame,
    strategy_run_summary_frame,
)


class ReportCatalogTests(unittest.TestCase):
    def test_collects_report_artifacts_and_cleans_only_junk(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            day = root / "2026-06-03"
            data_quality = day / "DataQuality"
            cross_validation = day / "CrossValidation"
            experiment = day / "Experiments" / "scan"
            data_quality.mkdir(parents=True)
            cross_validation.mkdir(parents=True)
            experiment.mkdir(parents=True)
            (day / "BacktestReport_03062026.docx").write_bytes(b"docx")
            (day / "ExperimentResearchReport_03062026.docx").write_bytes(b"docx")
            (day / "StrategyReviewReport_03062026.docx").write_bytes(b"docx")
            (day / "ReservedEmptyReport_03062026.docx").touch()
            (day / "RunMetadata_03062026.json").write_text("{}", encoding="utf-8")
            (data_quality / "DataQuality_sample_US_AAPL.json").write_text("{}", encoding="utf-8")
            (cross_validation / "CrossValidation_US_AAPL_yahoo_polygon.json").write_text("{}", encoding="utf-8")
            (experiment / "summary.csv").write_text(
                "run_id,param_short_window,param_long_window,total_return,sharpe,max_drawdown,cost_total\nrun1,5,30,0.1,1.2,-0.05,12.3\n",
                encoding="utf-8",
            )
            (experiment / "runs.json").write_text('[{"run_id": "run1"}]', encoding="utf-8")
            (experiment / "train_test_validation.json").write_text('{"validation_status": "Pass"}', encoding="utf-8")
            (experiment / "walk_forward_validation.json").write_text('{"validation_status": "Watch", "window_count": 2}', encoding="utf-8")
            (day / "OldReport.html").write_text("<html></html>", encoding="utf-8")
            (day / ".DS_Store").write_text("", encoding="utf-8")

            artifacts = collect_report_artifacts(root)
            counts = artifact_counts(root)
            experiments = experiment_summaries_frame(root)
            dry_run_targets = cleanup_report_junk(root, dry_run=True)
            removed = cleanup_report_junk(root, dry_run=False)

            self.assertEqual(len(artifacts), 9)
            self.assertNotIn("ReservedEmptyReport_03062026.docx", {artifact.name for artifact in artifacts})
            self.assertEqual(counts["Word Report"], 3)
            self.assertEqual(counts["Backtest Word Report"], 1)
            self.assertEqual(counts["Experiment Research Report"], 1)
            self.assertEqual(counts["Strategy Review Report"], 1)
            self.assertEqual(counts["Run Metadata"], 1)
            self.assertEqual(counts["Data Quality"], 1)
            self.assertEqual(counts["Cross Validation"], 1)
            self.assertEqual(counts["Experiment Summary"], 1)
            self.assertEqual(counts["Experiment Validation"], 1)
            self.assertEqual(counts["Walk Forward"], 1)
            self.assertEqual(len(experiments), 1)
            self.assertEqual(len(dry_run_targets), 3)
            self.assertEqual(len(removed), 3)
            self.assertTrue((day / "BacktestReport_03062026.docx").exists())
            self.assertTrue((day / "StrategyReviewReport_03062026.docx").exists())
            self.assertTrue((day / "RunMetadata_03062026.json").exists())
            self.assertFalse((day / "ReservedEmptyReport_03062026.docx").exists())
            self.assertFalse((day / "OldReport.html").exists())
            self.assertFalse((day / ".DS_Store").exists())

    def test_load_experiment_detail_reads_best_run_and_params(self):
        with TemporaryDirectory() as tmp:
            experiment = Path(tmp) / "Experiments" / "scan"
            experiment.mkdir(parents=True)
            summary_path = experiment / "summary.csv"
            summary_path.write_text(
                "run_id,param_short_window,param_long_window,total_return,sharpe,max_drawdown,cost_total\nbest,5,30,0.2,1.5,-0.04,10.0\nother,10,60,0.1,0.9,-0.08,15.0\n",
                encoding="utf-8",
            )
            (experiment / "runs.json").write_text('[{"run_id": "best"}, {"run_id": "other"}]', encoding="utf-8")
            (experiment / "train_test_validation.json").write_text('{"validation_status": "Pass", "generalization_ratio": 0.8}', encoding="utf-8")
            (experiment / "walk_forward_validation.json").write_text('{"validation_status": "Watch", "window_count": 3}', encoding="utf-8")

            detail = load_experiment_detail(summary_path)

            self.assertEqual(detail["experiment"], "scan")
            self.assertEqual(detail["run_count"], 2)
            self.assertEqual(detail["best_run"]["run_id"], "best")
            self.assertEqual(detail["best_params"]["short_window"], 5)
            self.assertIn("max_drawdown", detail["metric_columns"])
            self.assertEqual(len(detail["runs"]), 2)
            self.assertIn("stability", detail)
            self.assertIn("stability_status", detail["stability"])
            self.assertEqual(detail["train_test_validation"]["validation_status"], "Pass")
            self.assertEqual(detail["walk_forward_validation"]["validation_status"], "Watch")
            self.assertIn("risk_gate", detail)
            self.assertIn(detail["risk_gate"]["status"], {"ContinueResearch", "WatchOnly", "NeedsMoreEvidence", "DoNotUse"})

    def test_filters_report_artifacts_by_selected_types(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            day = root / "2026-06-03"
            other_day = root / "2026-06-04"
            day.mkdir(parents=True)
            other_day.mkdir(parents=True)
            (day / "BacktestReport_03062026.docx").write_bytes(b"docx")
            (day / "ExperimentResearchReport_03062026.docx").write_bytes(b"docx")
            (day / "StrategyReviewReport_03062026.docx").write_bytes(b"docx")
            (day / "RunMetadata_03062026.json").write_text("{}", encoding="utf-8")
            (other_day / "StrategyReviewReport_04062026.docx").write_bytes(b"docx")

            artifacts = report_artifacts_frame(root)
            strategy_reviews = filter_report_artifacts_frame(artifacts, ["Strategy Review Report"])
            strategy_reviews_for_day = filter_report_artifacts_frame(artifacts, ["Strategy Review Report"], ["2026-06-03"])
            word_reports = filter_report_artifacts_frame(
                artifacts,
                ["Backtest Word Report", "Experiment Research Report", "Strategy Review Report"],
            )
            empty_selection = filter_report_artifacts_frame(artifacts, [])

            self.assertEqual(len(strategy_reviews), 2)
            self.assertEqual(set(strategy_reviews["artifact_type"]), {"Strategy Review Report"})
            self.assertEqual(len(strategy_reviews_for_day), 1)
            self.assertEqual(strategy_reviews_for_day.iloc[0]["date_folder"], "2026-06-03")
            self.assertEqual(len(word_reports), 4)
            self.assertTrue(empty_selection.empty)

    def test_searches_report_artifacts_and_finds_latest_word_report(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            day = root / "2026-06-03"
            other_day = root / "2026-06-04"
            day.mkdir(parents=True)
            other_day.mkdir(parents=True)
            (day / "BacktestReport_03062026.docx").write_bytes(b"docx")
            (day / "RunMetadata_03062026.json").write_text("{}", encoding="utf-8")
            (other_day / "StrategyReviewReport_04062026.docx").write_bytes(b"docx")

            artifacts = report_artifacts_frame(root)
            by_name = search_report_artifacts_frame(artifacts, "strategyreview")
            by_type = search_report_artifacts_frame(artifacts, "run metadata")
            by_date = search_report_artifacts_frame(artifacts, "2026-06-03")
            by_path = search_report_artifacts_frame(artifacts, "04062026")
            no_query = search_report_artifacts_frame(artifacts, " ")
            latest = latest_report_artifact(artifacts, WORD_REPORT_TYPES)

            self.assertEqual(len(by_name), 1)
            self.assertEqual(by_name.iloc[0]["artifact_type"], "Strategy Review Report")
            self.assertEqual(len(by_type), 1)
            self.assertEqual(by_type.iloc[0]["artifact_type"], "Run Metadata")
            self.assertEqual(len(by_date), 2)
            self.assertEqual(len(by_path), 1)
            self.assertEqual(len(no_query), len(artifacts))
            self.assertIsNotNone(latest)
            self.assertIn(latest["artifact_type"], WORD_REPORT_TYPES)

    def test_run_metadata_summaries_classify_status(self):
        with TemporaryDirectory() as tmp:
            day = Path(tmp) / "2026-06-03"
            day.mkdir(parents=True)
            (day / "RunMetadataPass_03062026.json").write_text(
                '{"metrics": {"total_return": 0.2, "annualized_return": 0.18, "sharpe": 1.2, "max_drawdown": -0.08, "cost_total": 100, "ending_equity": 120000, "trade_count": 10}, "metadata": {"strategy": {"strategy_id": "ma_crossover"}, "backtest": {"initial_cash": 100000}}, "decision_quality": {"status": "NeedsMoreEvidence", "score": 64, "missing_evidence": ["多源交叉校验"]}}',
                encoding="utf-8",
            )
            (day / "RunMetadataReview_03062026.json").write_text(
                '{"metrics": {"total_return": -0.1, "annualized_return": -0.08, "sharpe": -0.2, "max_drawdown": -0.3, "cost_total": 100, "ending_equity": 90000, "trade_count": 8}, "metadata": {"strategy": {"strategy_id": "ma_crossover"}, "backtest": {"initial_cash": 100000}}}',
                encoding="utf-8",
            )

            frame = run_metadata_summaries_frame(tmp)

            self.assertEqual(len(frame), 2)
            self.assertEqual(set(frame["status"]), {"Pass", "Review"})
            self.assertIn("research_status", frame.columns)
            self.assertIn("decision_quality_score", frame.columns)
            self.assertIn("missing_evidence_count", frame.columns)
            pass_row = frame[frame["status"] == "Pass"].iloc[0]
            self.assertEqual(pass_row["research_status"], "NeedsMoreEvidence")
            self.assertEqual(pass_row["decision_quality_score"], 64)
            self.assertEqual(pass_row["missing_evidence_count"], 1)
            self.assertIn("metadata_path", frame.columns)

    def test_report_dashboard_frames_summarize_assets_runs_and_strategies(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            day = root / "2026-06-03"
            other_day = root / "2026-06-04"
            data_quality = day / "DataQuality"
            experiment = other_day / "Experiments" / "scan"
            data_quality.mkdir(parents=True)
            experiment.mkdir(parents=True)
            (day / "BacktestReport_03062026.docx").write_bytes(b"docx")
            (day / "RunMetadataPass_03062026.json").write_text(
                '{"metrics": {"total_return": 0.2, "annualized_return": 0.18, "sharpe": 1.2, "max_drawdown": -0.08, "cost_total": 100, "ending_equity": 120000, "trade_count": 10}, "metadata": {"strategy": {"strategy_id": "ma_crossover"}, "backtest": {"initial_cash": 100000}}}',
                encoding="utf-8",
            )
            (other_day / "RunMetadataReview_04062026.json").write_text(
                '{"metrics": {"total_return": -0.1, "annualized_return": -0.08, "sharpe": -0.2, "max_drawdown": -0.3, "cost_total": 300, "ending_equity": 90000, "trade_count": 8}, "metadata": {"strategy": {"strategy_id": "alipay"}, "backtest": {"initial_cash": 100000}}}',
                encoding="utf-8",
            )
            (data_quality / "DataQuality_sample_US_AAPL.json").write_text("{}", encoding="utf-8")
            (experiment / "summary.csv").write_text(
                "run_id,param_short_window,total_return,sharpe,max_drawdown,cost_total\nbest,5,0.15,1.1,-0.05,12.3\n",
                encoding="utf-8",
            )

            artifacts = report_artifacts_frame(root)
            counts = artifact_counts(root)
            runs = run_metadata_summaries_frame(root)
            experiments = experiment_summaries_frame(root)
            activity = report_activity_frame(artifacts)
            status_counts = run_status_counts_frame(runs)
            strategy_summary = strategy_run_summary_frame(runs)
            cards = report_dashboard_cards(counts, artifacts, runs, experiments, date_folder_count=2)

            self.assertEqual(activity["total_artifacts"].sum(), len(artifacts))
            self.assertIn("Review", set(status_counts["status"]))
            self.assertEqual(int(status_counts.loc[status_counts["status"] == "Review", "count"].iloc[0]), 1)
            self.assertEqual(set(strategy_summary["strategy_id"]), {"ma_crossover", "alipay"})
            self.assertEqual(int(strategy_summary["run_count"].sum()), 2)
            self.assertEqual(cards[0]["value"], counts["Word Report"])
            self.assertEqual(cards[2]["value"], 1)

    def test_export_experiment_docx_contains_validation_and_risk_sections(self):
        with TemporaryDirectory() as tmp:
            experiment = Path(tmp) / "Experiments" / "scan"
            experiment.mkdir(parents=True)
            summary_path = experiment / "summary.csv"
            summary_path.write_text(
                "run_id,strategy_id,param_short_window,param_long_window,total_return,annualized_return,sharpe,max_drawdown,win_rate,trade_count,cost_total,ending_equity\n"
                "best,ma_crossover,5,30,0.2,0.18,1.5,-0.04,0.6,10,10.0,120000\n"
                "other,ma_crossover,10,60,0.1,0.08,0.9,-0.08,0.5,8,15.0,110000\n",
                encoding="utf-8",
            )
            (experiment / "runs.json").write_text('[{"run_id": "best"}, {"run_id": "other"}]', encoding="utf-8")
            (experiment / "stability.json").write_text(
                '{"score_metric": "sharpe", "best_score": 1.5, "top_quantile_mean": 1.2, "top_quantile_threshold": 1.0, "neighbor_mean": 1.1, "neighbor_count": 2, "parameter_coverage": 0.5, "stability_status": "Stable", "notes": "Stable enough for research."}',
                encoding="utf-8",
            )
            (experiment / "train_test_validation.json").write_text(
                '{"validation_status": "Pass", "generalization_ratio": 0.8, "train_score": 1.5, "test_score": 1.2, "notes": "Pass."}',
                encoding="utf-8",
            )
            (experiment / "walk_forward_validation.json").write_text(
                '{"validation_status": "Watch", "window_count": 2, "pass_count": 1, "watch_count": 1, "failed_count": 0, "average_train_score": 1.4, "average_test_score": 0.9, "average_generalization_ratio": 0.64, "notes": "Watch.", "windows": [{"window": 1, "validation_status": "Pass", "train_score": 1.4, "test_score": 1.0, "best_params": {"short_window": 5}}]}',
                encoding="utf-8",
            )
            output_path = Path(tmp) / "ExperimentResearchReport_03062026.docx"

            report = export_experiment_docx(summary_path, output_path=output_path)

            self.assertTrue(report.exists())
            doc = Document(report)
            text = "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
            self.assertIn("QuantLab 实验研究报告", text)
            self.assertIn("Experiment Research Report", text)
            self.assertIn("研究风险闸门", text)
            self.assertIn("Research Risk Gate", text)
            self.assertIn("参数稳定性", text)
            self.assertIn("Parameter Stability", text)
            self.assertIn("样本内/样本外验证", text)
            self.assertIn("Train-Test Validation", text)
            self.assertIn("滚动样本外验证", text)
            self.assertIn("Walk-Forward Validation", text)
            self.assertIn("文件追溯", text)
            self.assertIn("总收益参数热力图", text)
            self.assertIn("夏普参数热力图", text)
            self.assertIn("最大回撤参数热力图", text)
            self.assertGreaterEqual(len(doc.inline_shapes), 5)


if __name__ == "__main__":
    unittest.main()
