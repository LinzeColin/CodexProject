import tempfile
import unittest
from pathlib import Path

from quantlab.approvals import StrategyApprovalRegistry
from quantlab.reports import export_strategy_review_docx
from quantlab.strategies import create_strategy_template, evaluate_strategy_code_quality, evaluate_strategy_readiness_gate, normalize_strategy_id, run_strategy_smoke_test, strategy_class_name, strategy_code_quality_rows, strategy_smoke_test_rows


class CandidateStub:
    def __init__(self, strategy_id, version, quality_status, approval_status="Pending", missing_items=()):
        self.strategy_id = strategy_id
        self.version = version
        self.quality_status = quality_status
        self.approval_status = approval_status
        self.missing_items = missing_items


class StrategyTemplateTests(unittest.TestCase):
    def test_normalizes_strategy_id_and_class_name(self):
        self.assertEqual(normalize_strategy_id(" Custom Momentum Filter "), "custom_momentum_filter")
        self.assertEqual(strategy_class_name("custom_momentum_filter"), "CustomMomentumFilterStrategy")
        self.assertEqual(strategy_class_name("demo_custom_strategy"), "DemoCustomStrategy")
        with self.assertRaises(ValueError):
            normalize_strategy_id("123_bad")

    def test_create_strategy_template_writes_draft_and_pending_approval(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            registry = StrategyApprovalRegistry(path=root / "StrategyApprovals.json")

            artifact = create_strategy_template(
                strategy_id="custom_momentum_filter",
                display_name="自定义动量过滤",
                display_name_en="Custom Momentum Filter",
                category="趋势跟随 Trend Following",
                return_source="行为偏差 Behavioral Bias",
                thesis="价格动量可能延续。",
                failure="震荡市场可能失效。",
                strategy_dir=root / "strategies",
                profile_dir=root / "profiles",
                approval_registry=registry,
            )

            strategy_file = Path(artifact.strategy_file)
            profile_file = Path(artifact.profile_file)
            self.assertTrue(strategy_file.exists())
            self.assertTrue(profile_file.exists())
            self.assertEqual(artifact.status, "Pending")
            self.assertEqual(len(registry.records()), 1)
            self.assertEqual(registry.records()[0].status, "Pending")
            strategy_text = strategy_file.read_text(encoding="utf-8")
            profile_text = profile_file.read_text(encoding="utf-8")
            self.assertIn("CustomMomentumFilterStrategy", strategy_text)
            self.assertIn("Research only, not approved by default", strategy_text)
            self.assertIn("审批状态", profile_text)
            self.assertIn("Approval Status", profile_text)
            quality = evaluate_strategy_code_quality(strategy_file)
            self.assertEqual(quality.status, "CodeDraft")
            self.assertIn("仍是空仓模板 Still flat draft template", quality.findings)
            smoke = run_strategy_smoke_test(strategy_file)
            self.assertEqual(smoke.status, "SmokePass")
            self.assertGreater(smoke.rows, 0)

    def test_create_strategy_template_does_not_overwrite_by_default(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            registry = StrategyApprovalRegistry(path=root / "StrategyApprovals.json")
            kwargs = {
                "strategy_id": "custom_signal",
                "display_name": "自定义信号",
                "display_name_en": "Custom Signal",
                "category": "研究 Research",
                "return_source": "行为偏差 Behavioral Bias",
                "thesis": "测试假设。",
                "failure": "测试失效环境。",
                "strategy_dir": root / "strategies",
                "profile_dir": root / "profiles",
                "approval_registry": registry,
            }
            create_strategy_template(**kwargs)
            with self.assertRaises(FileExistsError):
                create_strategy_template(**kwargs)

    def test_strategy_code_quality_marks_non_template_code_ready(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            strategy_file = root / "custom_ready.py"
            strategy_file.write_text(
                '''
from __future__ import annotations

import pandas as pd

from quantlab.strategies.base import Strategy, StrategyResult, finalize_signal_frame


class CustomReadyStrategy(Strategy):
    strategy_id = "custom_ready"
    version = "0.1.0"
    description = "Ready candidate."

    def __init__(self, lookback: int = 20):
        if lookback < 2:
            raise ValueError("lookback must be >= 2")
        super().__init__(lookback=lookback)
        self.lookback = lookback

    def generate_signals(self, data: pd.DataFrame) -> StrategyResult:
        momentum = data["close"].pct_change(self.lookback)
        target = pd.Series(0.0, index=data.index)
        target[momentum > 0] = 1.0
        signals = finalize_signal_frame(data, target)
        return StrategyResult(signals=signals, metadata=self.metadata())
''',
                encoding="utf-8",
            )

            quality = evaluate_strategy_code_quality(strategy_file)
            rows = strategy_code_quality_rows(root)

            self.assertEqual(quality.status, "CodeReadyForReview")
            self.assertEqual(quality.score, 100)
            self.assertEqual(quality.findings, ())
            self.assertEqual(rows[0]["策略编号 Strategy Id"], "custom_ready")
            self.assertEqual(rows[0]["代码状态 Code Status"], "CodeReadyForReview")
            smoke = run_strategy_smoke_test(strategy_file)
            smoke_rows = strategy_smoke_test_rows(root)
            self.assertEqual(smoke.status, "SmokePass")
            self.assertGreater(smoke.rows, 0)
            self.assertEqual(smoke_rows[0]["烟雾测试 Smoke Test"], "SmokePass")

    def test_strategy_smoke_test_fails_invalid_weights(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            strategy_file = root / "bad_weights.py"
            strategy_file.write_text(
                '''
from __future__ import annotations
import pandas as pd
from quantlab.strategies.base import Strategy, StrategyResult, finalize_signal_frame
class BadWeightsStrategy(Strategy):
    strategy_id = "bad_weights"
    version = "0.1.0"
    def generate_signals(self, data: pd.DataFrame) -> StrategyResult:
        signals = finalize_signal_frame(data, pd.Series(0.0, index=data.index))
        signals["target_weight"] = 2.0
        return StrategyResult(signals=signals, metadata=self.metadata())
''',
                encoding="utf-8",
            )

            smoke = run_strategy_smoke_test(strategy_file)

            self.assertEqual(smoke.status, "SmokeFail")
            self.assertTrue(any("outside" in finding for finding in smoke.findings))

    def test_strategy_readiness_gate_combines_profile_code_and_approval(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            strategy_file = root / "custom_ready.py"
            strategy_file.write_text(
                '''
from __future__ import annotations
import pandas as pd
from quantlab.strategies.base import Strategy, StrategyResult, finalize_signal_frame
class CustomReadyStrategy(Strategy):
    strategy_id = "custom_ready"
    version = "0.1.0"
    def __init__(self, lookback: int = 20):
        if lookback < 2:
            raise ValueError("lookback must be >= 2")
        super().__init__(lookback=lookback)
        self.lookback = lookback
    def generate_signals(self, data: pd.DataFrame) -> StrategyResult:
        target = pd.Series(0.0, index=data.index)
        target[data["close"].pct_change(self.lookback) > 0] = 1.0
        signals = finalize_signal_frame(data, target)
        return StrategyResult(signals=signals, metadata=self.metadata())
''',
                encoding="utf-8",
            )
            code_report = evaluate_strategy_code_quality(strategy_file)
            smoke_report = run_strategy_smoke_test(strategy_file)
            registry = StrategyApprovalRegistry(path=root / "StrategyApprovals.json")
            candidate = CandidateStub("custom_ready", "0.1.0", "ReadyForReview")

            missing_smoke = evaluate_strategy_readiness_gate(candidate, code_report, registry.records())
            not_ready = evaluate_strategy_readiness_gate(CandidateStub("custom_ready", "0.1.0", "Incomplete", missing_items=("参数设置 Parameter Settings",)), code_report, registry.records(), smoke_report=smoke_report)
            ready = evaluate_strategy_readiness_gate(candidate, code_report, registry.records(), smoke_report=smoke_report)
            registry.request_approval("custom_ready", "0.1.0", "Ready candidate")
            registry.approve("custom_ready", "0.1.0")
            approved = evaluate_strategy_readiness_gate(candidate, code_report, registry.records(), smoke_report=smoke_report)

            self.assertEqual(missing_smoke.status, "NotReady")
            self.assertEqual(not_ready.status, "NotReady")
            self.assertEqual(ready.status, "ReadyForReview")
            self.assertEqual(approved.status, "ApprovedForResearch")

    def test_strategy_readiness_gate_blocks_smoke_failure(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            strategy_file = root / "bad_weights.py"
            strategy_file.write_text(
                '''
from __future__ import annotations
import pandas as pd
from quantlab.strategies.base import Strategy, StrategyResult, finalize_signal_frame
class BadWeightsStrategy(Strategy):
    strategy_id = "bad_weights"
    version = "0.1.0"
    def __init__(self, lookback: int = 20):
        if lookback < 2:
            raise ValueError("lookback must be >= 2")
        super().__init__(lookback=lookback)
    def generate_signals(self, data: pd.DataFrame) -> StrategyResult:
        signals = finalize_signal_frame(data, pd.Series(0.0, index=data.index))
        signals["target_weight"] = 2.0
        return StrategyResult(signals=signals, metadata=self.metadata())
''',
                encoding="utf-8",
            )
            candidate = CandidateStub("bad_weights", "0.1.0", "ReadyForReview", approval_status="Approved")
            code_report = evaluate_strategy_code_quality(strategy_file)
            smoke_report = run_strategy_smoke_test(strategy_file)

            gate = evaluate_strategy_readiness_gate(candidate, code_report, [], smoke_report=smoke_report)

            self.assertEqual(smoke_report.status, "SmokeFail")
            self.assertEqual(gate.status, "NotReady")
            self.assertTrue(any("Smoke test" in reason for reason in gate.reasons))

    def test_export_strategy_review_docx_contains_gate_and_checks(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            strategy_dir = root / "strategies"
            profile_dir = root / "profiles"
            registry = StrategyApprovalRegistry(path=root / "StrategyApprovals.json")
            artifact = create_strategy_template(
                strategy_id="custom_review",
                display_name="自定义审查",
                display_name_en="Custom Review",
                category="研究 Research",
                return_source="行为偏差 Behavioral Bias",
                thesis="价格动量可能延续。",
                failure="震荡市场可能失效。",
                strategy_dir=strategy_dir,
                profile_dir=profile_dir,
                approval_registry=registry,
            )
            output = root / "StrategyReviewReport_03062026.docx"

            report = export_strategy_review_docx(artifact.profile_file, output_path=output)

            self.assertTrue(report.exists())
            doc = Document(report)
            text = "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
            self.assertIn("候选策略审查报告", text)
            self.assertIn("Candidate Strategy Review Report", text)
            self.assertIn("审批前综合门禁", text)
            self.assertIn("Pre-Approval Readiness Gate", text)
            self.assertIn("代码质量检查", text)
            self.assertIn("Code Quality Check", text)
            self.assertIn("烟雾测试", text)
            self.assertIn("Smoke Test", text)
            self.assertIn("审批记录", text)
            self.assertIn("Approval Records", text)


if __name__ == "__main__":
    unittest.main()
