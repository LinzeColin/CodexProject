import json
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from quantlab.backtest import BacktestConfig, BacktestEngine
from quantlab.config import REPORT_ROOT_DIR
from quantlab.data.quality import assess_bars
from quantlab.data.models import BarDataRequest
from quantlab.data.providers import SampleDataProvider
from quantlab.data.validation import CrossSourceValidationResult
from quantlab.reports import export_backtest_docx, report_filename, unique_report_path
from quantlab.reports.export import (
    _buy_and_hold_metrics,
    _monthly_return_matrix,
    report_evidence_summary,
    _report_bootstrap_simulations,
    _return_comparison_frame,
    _return_source_rows,
    _rolling_risk_frame,
    _strategy_review_rows,
)
from quantlab.strategies import MovingAverageCrossoverStrategy


class ReportTests(unittest.TestCase):
    def tearDown(self):
        for path in getattr(self, "_created_paths", []):
            path.unlink(missing_ok=True)

    def _track(self, *paths):
        self._created_paths = getattr(self, "_created_paths", [])
        self._created_paths.extend(paths)

    def test_report_exports_to_dated_codex_report_directory(self):
        data = SampleDataProvider(seed=9).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        path = export_backtest_docx(result)
        self.assertTrue(path.exists())
        self.assertEqual(path.parent.parent, REPORT_ROOT_DIR)
        self.assertTrue(path.name.startswith("BacktestReport"))
        self.assertIn("_", path.name)
        self.assertEqual(path.suffix, ".docx")
        metadata = path.with_name(path.stem.replace("BacktestReport", "RunMetadata", 1) + ".json")
        self.assertTrue(metadata.exists())
        payload = json.loads(metadata.read_text(encoding="utf-8"))
        self.assertIn("risk_gate", payload)
        self.assertIn("decision_quality", payload)
        self._track(path, metadata)

    def test_report_contains_strategy_and_bilingual_metrics(self):
        data = SampleDataProvider(seed=11).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        path = export_backtest_docx(result, report_name="ContentCheckReport")
        metadata = path.with_name(path.stem.replace("ContentCheckReport", "RunMetadata", 1) + ".json")
        self._track(path, metadata)
        doc = Document(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        combined = text + "\n" + table_text
        self.assertIn("策略说明", combined)
        self.assertIn("Strategy Description", combined)
        self.assertIn("总收益", combined)
        self.assertIn("Total Return", combined)
        self.assertIn("买入持有总收益", combined)
        self.assertIn("Buy And Hold Total Return", combined)
        self.assertIn("买入持有最大回撤", combined)
        self.assertIn("Buy And Hold Max Drawdown", combined)
        self.assertIn("结果判读", combined)
        self.assertIn("Result Interpretation", combined)
        self.assertIn("相对买入持有", combined)
        self.assertIn("Versus Buy And Hold", combined)
        self.assertIn("交易摩擦", combined)
        self.assertIn("Trading Friction", combined)
        self.assertIn("决策质量摘要", combined)
        self.assertIn("Decision Quality Score", combined)
        self.assertIn("报告证据层", combined)
        self.assertIn("Report Evidence Layer", combined)
        self.assertIn("关键缺失证据", combined)

    def test_report_evidence_summary_downgrades_when_critical_evidence_is_missing(self):
        data = SampleDataProvider(seed=12).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )

        evidence = report_evidence_summary(result)

        self.assertEqual(evidence["schema"], "QuantLabReportEvidenceV1")
        self.assertEqual(evidence["evidence_status"], "NeedsMoreEvidence")
        self.assertEqual(evidence["workflow"]["lineage_status"], "ManualOrLocalOnly")
        self.assertIn("数据质量报告", evidence["missing_evidence"])
        self.assertIn("多源交叉校验", evidence["missing_evidence"])
        self.assertIn("实体注册状态", evidence["missing_evidence"])

    def test_report_evidence_is_written_to_docx_and_run_metadata(self):
        data = SampleDataProvider(seed=14).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        quality = assess_bars(data, provider="sample", symbol="AAPL", market="US", interval="1d")
        cross = CrossSourceValidationResult(
            symbol="AAPL",
            market="US",
            interval="1d",
            providers=("sample", "yahoo_finance"),
            overlap_rows=5,
            max_close_diff_pct=0.0,
            mean_close_diff_pct=0.0,
            status="Pass",
            details=data.head(0),
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        result.metadata["entity"] = {"status": "TradableSymbol", "canonical_symbol": "AAPL", "market": "US"}
        result.metadata["workflow"] = {
            "workflow_input_id": "chatInput_test",
            "linked_request_id": "busRequest_test",
            "source_system": "UnitTest",
        }

        path = export_backtest_docx(
            result,
            report_name="EvidenceLayerReport",
            data_quality_report=quality,
            cross_validation_result=cross,
        )
        metadata = path.with_name(path.stem.replace("EvidenceLayerReport", "RunMetadata", 1) + ".json")
        self._track(path, metadata)
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        payload = json.loads(metadata.read_text(encoding="utf-8"))

        self.assertIn("报告证据层", text)
        self.assertIn("TradableSymbol", text)
        self.assertIn("chatInput_test", text)
        self.assertEqual(payload["report_evidence"]["entity"]["entity_status"], "TradableSymbol")
        self.assertEqual(payload["report_evidence"]["workflow"]["lineage_status"], "Linked")

    def test_report_contains_charts_and_quality_summaries(self):
        data = SampleDataProvider(seed=13).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        quality = assess_bars(data, provider="sample", symbol="AAPL", market="US", interval="1d")
        cross = CrossSourceValidationResult(
            symbol="AAPL",
            market="US",
            interval="1d",
            providers=("sample", "yahoo_finance"),
            overlap_rows=5,
            max_close_diff_pct=0.0123,
            mean_close_diff_pct=0.0045,
            status="Review",
            details=data.head(0),
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        path = export_backtest_docx(result, report_name="ChartQualityReport", data_quality_report=quality, cross_validation_result=cross)
        metadata = path.with_name(path.stem.replace("ChartQualityReport", "RunMetadata", 1) + ".json")
        self._track(path, metadata)
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        self.assertGreaterEqual(len(doc.inline_shapes), 10)
        self.assertIn("数据质量摘要", text)
        self.assertIn("Data Quality Summary", text)
        self.assertIn("多源交叉校验摘要", text)
        self.assertIn("Cross-Source Validation Summary", text)
        self.assertIn("策略审批状态", text)
        self.assertIn("策略诊断", text)
        self.assertIn("Strategy Diagnostics", text)
        self.assertIn("交易质量", text)
        self.assertIn("Trade Quality", text)
        self.assertIn("成本压力", text)
        self.assertIn("Cost Stress", text)
        self.assertIn("市场环境分层", text)
        self.assertIn("Market Regime Breakdown", text)
        self.assertIn("研究风险闸门", text)
        self.assertIn("Research Risk Gate", text)
        self.assertIn("决策质量摘要", text)
        self.assertIn("评分维度", text)
        self.assertIn("历史模拟暴露统计", text)
        self.assertIn("研究状态", text)
        self.assertIn("Research Status", text)
        self.assertIn("风险分数", text)
        self.assertIn("Risk Score", text)
        self.assertIn("触发原因", text)
        self.assertIn("Triggered Reasons", text)
        self.assertIn("建议动作", text)
        self.assertIn("Suggested Actions", text)
        self.assertIn("策略收益与买入持有收益对比", text)
        self.assertIn("策略、目标与相对收益曲线", text)
        self.assertIn("月度收益热力图", text)
        self.assertIn("滚动夏普与波动率", text)
        self.assertIn("Bootstrap 鲁棒性验证", text)
        self.assertIn("模拟总收益分布", text)
        self.assertIn("模拟最大回撤分布", text)
        self.assertIn("亏损概率", text)
        self.assertIn("策略研究审查", text)
        self.assertIn("我赚的是什么钱？", text)
        self.assertIn("扣除费用后是否仍有效？", text)
        self.assertIn("失效后系统如何停止交易？", text)
        self.assertIn("策略收益来源", text)
        self.assertIn("风险溢价", text)
        self.assertIn("Risk Premium", text)
        self.assertIn("行为偏差", text)
        self.assertIn("Behavioral Bias", text)
        self.assertIn("组合优势", text)
        self.assertIn("Portfolio Advantage", text)

    def test_report_buy_and_hold_metrics_include_max_drawdown(self):
        data = SampleDataProvider(seed=21).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-12-31")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )

        metrics = _buy_and_hold_metrics(result)

        self.assertIn("buy_hold_max_drawdown", metrics)
        self.assertIn("buy_hold_sharpe", metrics)
        self.assertLessEqual(metrics["buy_hold_max_drawdown"], 0.0)

    def test_report_visualization_frames_are_research_ready(self):
        data = SampleDataProvider(seed=16).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2021-12-31")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        price_frame = result.positions[["datetime", "close"]].drop_duplicates("datetime")

        comparison = _return_comparison_frame(result.equity_curve, price_frame)
        monthly = _monthly_return_matrix(result.equity_curve)
        rolling = _rolling_risk_frame(result.equity_curve)

        self.assertIn("relative_return", comparison.columns)
        self.assertFalse(comparison.empty)
        self.assertFalse(monthly.empty)
        self.assertFalse(rolling.empty)
        self.assertIn("rolling_sharpe", rolling.columns)

    def test_report_review_uses_custom_strategy_spec_when_present(self):
        custom_spec = {
            "display_name": "自定义均值回归",
            "display_name_en": "Custom Mean Reversion",
            "logic_key": "mean_reversion",
            "indicator_keys": ["rsi", "bollinger"],
            "return_source": "行为偏差 Behavioral Bias；执行优势 Execution Advantage",
            "return_source_en": "Behavioral Bias; Execution Advantage",
            "thesis": "短期过度偏离后可能回归。",
            "thesis_en": "Short-term dislocation may revert.",
            "failure": "单边趋势可能失效。",
            "failure_en": "One-way trends may invalidate it.",
        }
        strategy = {"strategy_id": "custom_mean_reversion", "custom_strategy_spec": custom_spec}
        result = SimpleNamespace(
            metrics={
                "total_return": 0.12,
                "max_drawdown": -0.08,
                "sharpe": 1.1,
                "trade_count": 4,
                "cost_total": 100.0,
                "ending_equity": 112000.0,
            }
        )

        review_rows = _strategy_review_rows(result, strategy, "Approved")
        source_rows = _return_source_rows(strategy)

        self.assertIn("短期过度偏离", review_rows[1][2])
        self.assertIn("单边趋势", review_rows[5][2])
        self.assertTrue(any(row[0] == "行为偏差" and "Primary" in row[2] for row in source_rows))
        self.assertTrue(any(row[0] == "执行优势" and "Primary" in row[2] for row in source_rows))

    def test_report_filename_uses_ddmmyyyy(self):
        from datetime import datetime

        self.assertEqual(report_filename("BacktestReport", day=datetime(2026, 6, 3)), "BacktestReport_03062026.docx")

    def test_report_bootstrap_simulation_count_has_10000_floor(self):
        import os

        original = os.environ.pop("QUANTLAB_TEST_BOOTSTRAP_SIMULATIONS", None)
        try:
            self.assertEqual(_report_bootstrap_simulations(), 10000)
            os.environ["QUANTLAB_TEST_BOOTSTRAP_SIMULATIONS"] = "300"
            self.assertEqual(_report_bootstrap_simulations(), 10000)
        finally:
            if original is not None:
                os.environ["QUANTLAB_TEST_BOOTSTRAP_SIMULATIONS"] = original
            else:
                os.environ.pop("QUANTLAB_TEST_BOOTSTRAP_SIMULATIONS", None)

    def test_report_export_does_not_overwrite_same_day_report(self):
        data = SampleDataProvider(seed=10).get_bars(
            BarDataRequest(symbol="AAPL", market="US", interval="1d", start="2020-01-01", end="2020-06-30")
        )
        result = BacktestEngine(BacktestConfig(initial_cash=100_000)).run(
            data, MovingAverageCrossoverStrategy(short_window=5, long_window=20)
        )
        first = export_backtest_docx(result, report_name="TestBacktestReport")
        second = export_backtest_docx(result, report_name="TestBacktestReport")
        self.assertNotEqual(first, second)
        self.assertTrue(first.exists())
        self.assertTrue(second.exists())
        self.assertIn("_", second.name)
        first_metadata = first.with_name(first.stem.replace("TestBacktestReport", "RunMetadata", 1) + ".json")
        second_metadata = second.with_name(second.stem.replace("TestBacktestReport", "RunMetadata", 1) + ".json")
        self._track(first, second, first_metadata, second_metadata)

    def test_unique_report_path_reserves_paths_to_avoid_name_races(self):
        first = unique_report_path("ReservedReportRaceCheck")
        second = unique_report_path("ReservedReportRaceCheck")

        self.assertNotEqual(first, second)
        self.assertTrue(first.exists())
        self.assertTrue(second.exists())
        self.assertEqual(first.stat().st_size, 0)
        self.assertEqual(second.stat().st_size, 0)
        self._track(first, second)


if __name__ == "__main__":
    unittest.main()
